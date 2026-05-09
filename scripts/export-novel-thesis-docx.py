from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BOOK_ROOT = Path("test-artifacts/novel-book")
CHAPTER_RE = re.compile(r"^#\s+(.+?)\s*$")


def set_run_font(run, east_asia: str, ascii_font: str | None = None, size: Pt | None = None, bold: bool | None = None):
    font = run.font
    font.name = ascii_font or east_asia
    if size is not None:
        font.size = size
    if bold is not None:
        font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font or "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), ascii_font or "Times New Roman")


def set_style_font(style, east_asia: str, ascii_font: str, size_pt: float, bold: bool = False):
    font = style.font
    font.name = ascii_font
    font.size = Pt(size_pt)
    font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)


def add_field(paragraph, instruction: str, placeholder: str = ""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instr_run._r.append(instr)

    sep_run = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(sep)

    if placeholder:
        paragraph.add_run(placeholder)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, "PAGE", "1")
    for run in paragraph.runs:
        set_run_font(run, "宋体", "Times New Roman", Pt(10.5))


def set_section_layout(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)


def configure_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, "宋体", "Times New Roman", 12)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(20)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    heading1 = styles["Heading 1"]
    set_style_font(heading1, "黑体", "Times New Roman", 16, True)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.first_line_indent = Pt(0)
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading1.paragraph_format.line_spacing = Pt(24)

    heading2 = styles["Heading 2"]
    set_style_font(heading2, "黑体", "Times New Roman", 14, True)
    heading2.paragraph_format.first_line_indent = Pt(0)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading2.paragraph_format.line_spacing = Pt(22)

    if "Thesis Body" not in styles:
        body = styles.add_style("Thesis Body", 1)
    else:
        body = styles["Thesis Body"]
    body.base_style = normal
    set_style_font(body, "宋体", "Times New Roman", 12)
    body.paragraph_format.first_line_indent = Pt(24)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    body.paragraph_format.line_spacing = Pt(20)
    body.paragraph_format.space_after = Pt(0)


def add_centered_line(doc: Document, text: str, font: str = "宋体", size: float = 12, bold: bool = False, after: float = 0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, font, "Times New Roman", Pt(size), bold)
    return p


def add_label_line(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label}：{value}")
    set_run_font(run, "宋体", "Times New Roman", Pt(14))


def add_thesis_heading(doc: Document, text: str, level: int = 1):
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Pt(0)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(24)
        run = p.add_run(text)
        set_run_font(run, "黑体", "Times New Roman", Pt(16), True)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(22)
        run = p.add_run(text)
        set_run_font(run, "黑体", "Times New Roman", Pt(14), True)
    return p


def add_thesis_body(doc: Document, text: str, first_line_indent: bool = True):
    p = doc.add_paragraph(style="Thesis Body")
    p.paragraph_format.first_line_indent = Pt(24 if first_line_indent else 0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", Pt(12))
    return p


def add_front_matter(doc: Document):
    add_centered_line(doc, "硕士学位论文", "黑体", 18, True, 24)
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    title_run = p.add_run("雨幕下的赤铜门")
    set_run_font(title_run, "黑体", "Times New Roman", Pt(22), True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    sub_run = p.add_run("原创现代幻想长篇文本")
    set_run_font(sub_run, "宋体", "Times New Roman", Pt(14))

    for _ in range(4):
        doc.add_paragraph()

    add_label_line(doc, "作者", "林知遥课题组文本整理")
    add_label_line(doc, "学科专业", "中国现当代文学")
    add_label_line(doc, "研究方向", "现代幻想叙事与城市记忆")
    add_label_line(doc, "指导教师", "唐听雨")
    add_label_line(doc, "提交日期", "2026年4月")

    doc.add_page_break()

    add_thesis_heading(doc, "摘 要", level=1)
    abstracts = [
        "《雨幕下的赤铜门》是一部原创现代都市幻想长篇文本。作品以架空海港城市临川港和隐秘学院白塔书院为叙事中心，围绕赤铜门、古相、灰潮和归潮协议展开，讲述少年林知遥在被命运选中之前学会拒绝、协商和承担边界的故事。",
        "全文以雨、铜、钟表、海潮、档案和影线为主要意象，通过林知遥、叶澜、沈既白、宋砚、唐听雨等人物的行动，讨论血统、教育、权力、记忆、牺牲与个体拒绝权之间的关系。文本最终将冲突从单一牺牲式封门转向可撤回同意、公开记录和分布式边界维护，形成“归潮协议”的伦理核心。",
    ]
    for text in abstracts:
        add_thesis_body(doc, text)
    add_thesis_body(doc, "关键词：临川港；赤铜门；古相；灰潮；边界；拒绝权", first_line_indent=False)

    doc.add_page_break()
    add_thesis_heading(doc, "Abstract", level=1)
    english = [
        "Rain Beneath the Red-Copper Gate is an original modern fantasy novel set in the fictional harbor city of Linchuan and the hidden White Tower academy. Through the Red-Copper Gate, ancient echoes, gray tide contamination, and the Returning Tide Protocol, the narrative follows a young protagonist who learns to refuse imposed destiny before learning to negotiate with it.",
        "The manuscript uses rain, copper, clocks, tides, archives, and shadows as recurring images. Its ethical focus is the rejection of systems that normalize sacrifice and the construction of boundaries based on consent, refusal, public record, and review.",
    ]
    for text in english:
        add_thesis_body(doc, text, first_line_indent=False)
    add_thesis_body(
        doc,
        "Key words: Linchuan Harbor; Red-Copper Gate; Ancient Echo; Gray Tide; Boundary; Right of Refusal",
        first_line_indent=False,
    )

    doc.add_page_break()
    add_thesis_heading(doc, "目 录", level=1)
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.first_line_indent = Pt(0)
    add_field(toc_p, 'TOC \\o "1-2" \\h \\z \\u', "右键更新域以生成目录")
    doc.add_page_break()


def read_chapter(path: Path):
    lines = path.read_text(encoding="utf-8").splitlines()
    title = None
    body: list[str] = []
    for line in lines:
        match = CHAPTER_RE.match(line)
        if match and title is None:
            title = match.group(1).strip()
            continue
        if line.strip():
            body.append(line.strip())
    return title or path.stem, body


def add_chapters(doc: Document, chapter_paths: list[Path]):
    for index, path in enumerate(chapter_paths, start=1):
        title, paragraphs = read_chapter(path)
        if index > 1:
            doc.add_page_break()
        add_thesis_heading(doc, title, level=1)
        for text in paragraphs:
            add_thesis_body(doc, text)


def add_back_matter(doc: Document):
    doc.add_page_break()
    add_thesis_heading(doc, "参考文献", level=1)
    add_thesis_body(doc, "本文为原创文学文本整理稿，正文创作未引用外部文献。")

    doc.add_page_break()
    add_thesis_heading(doc, "致 谢", level=1)
    add_thesis_body(
        doc,
        "感谢所有参与文本生成、设定整理、格式校验和文档排版的协作者。本稿以可审查、可继续修订的方式保存，便于后续插图、校对和正式装订。",
    )


def enable_update_fields(docx_path: Path):
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as package:
        settings_path = "word/settings.xml"
        xml = package.read(settings_path).decode("utf-8")
        if "w:updateFields" not in xml:
            xml = xml.replace(
                "</w:settings>",
                '<w:updateFields w:val="true"/></w:settings>',
            )
        else:
            return

        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as rewritten:
            for item in package.infolist():
                data = xml.encode("utf-8") if item.filename == settings_path else package.read(item.filename)
                rewritten.writestr(item, data)

    tmp_path.replace(docx_path)


def build_docx(book_root: Path, output: Path):
    doc = Document()
    set_section_layout(doc.sections[0])
    configure_styles(doc)
    add_page_number(doc.sections[0])

    add_front_matter(doc)
    chapter_paths = sorted((book_root / "chapters").glob("chapter-*.md"))
    if len(chapter_paths) != 96:
        raise RuntimeError(f"Expected 96 chapter files, found {len(chapter_paths)}")
    add_chapters(doc, chapter_paths)
    add_back_matter(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    enable_update_fields(output)
    return chapter_paths


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--book-root", default=str(BOOK_ROOT))
    parser.add_argument(
        "--out",
        default=str(BOOK_ROOT / "wps-thesis" / "yumuxia-de-chitongmen-master-thesis-format.docx"),
    )
    args = parser.parse_args()

    book_root = Path(args.book_root)
    output = Path(args.out)
    chapters = build_docx(book_root, output)
    print(f"DOCX={output.resolve()}")
    print(f"CHAPTERS={len(chapters)}")
    print(f"BYTES={output.stat().st_size}")


if __name__ == "__main__":
    main()
