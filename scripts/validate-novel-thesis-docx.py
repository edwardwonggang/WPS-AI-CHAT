from __future__ import annotations

import argparse
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


BOOK_ROOT = Path("test-artifacts/novel-book")
DEFAULT_DOCX = BOOK_ROOT / "wps-thesis" / "yumuxia-de-chitongmen-master-thesis-format-v2.docx"
CHAPTER_RE = re.compile(r"^#\s+(.+?)\s*$")
DOUBLE_CHAPTER_RE = re.compile(r"^(第[一二三四五六七八九十百千万零〇两]+章)\s+第[一二三四五六七八九十百千万零〇两]+章")


def fail(message: str):
    print(f"FAIL={message}")
    raise SystemExit(1)


def near(actual: float, expected: float, tolerance: float = 0.05) -> bool:
    return abs(actual - expected) <= tolerance


def style_east_asia_font(style) -> str | None:
    r_pr = style._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return None
    return r_pr.rFonts.get(qn("w:eastAsia"))


def style_ascii_font(style) -> str | None:
    r_pr = style._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return None
    return r_pr.rFonts.get(qn("w:ascii"))


def run_east_asia_font(run) -> str | None:
    r_pr = run._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return None
    return r_pr.rFonts.get(qn("w:eastAsia"))


def run_ascii_font(run) -> str | None:
    r_pr = run._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return None
    return r_pr.rFonts.get(qn("w:ascii"))


def paragraph_has_run_font(paragraph, east_asia: str, ascii_font: str | None = None) -> bool:
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        if run_east_asia_font(run) != east_asia:
            return False
        if ascii_font is not None and run_ascii_font(run) != ascii_font:
            return False
    return True


def chapter_title_and_body(path: Path):
    title = None
    body: list[str] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        match = CHAPTER_RE.match(line)
        if match and title is None:
            title = match.group(1).strip()
            continue
        if line.strip():
            body.append(line.strip())
    return title or path.stem, body


def validate_zip(docx_path: Path):
    update_fields = False
    with zipfile.ZipFile(docx_path, "r") as package:
        names = package.namelist()
        duplicates = [name for name, count in Counter(names).items() if count > 1]
        if duplicates:
            fail("duplicate zip entries: " + ", ".join(duplicates[:5]))

        settings = package.read("word/settings.xml").decode("utf-8", errors="replace")
        update_fields = "w:updateFields" in settings

        document_xml = package.read("word/document.xml").decode("utf-8", errors="replace")
        normalized_document_xml = document_xml.replace("&quot;", '"')
        if "TOC" not in normalized_document_xml or '\\o "1-2"' not in normalized_document_xml:
            fail("table-of-contents field was not found")

        footer_xml = "\n".join(
            package.read(name).decode("utf-8", errors="replace")
            for name in names
            if name.startswith("word/footer") and name.endswith(".xml")
        )
        if "PAGE" not in footer_xml:
            fail("page-number field was not found")

    return {"update_fields": update_fields}


def validate_layout_and_styles(doc: Document):
    section = doc.sections[0]
    checks = [
        (near(section.page_width.cm, 21.0), "page width is not A4 21.0 cm"),
        (near(section.page_height.cm, 29.7), "page height is not A4 29.7 cm"),
        (near(section.top_margin.cm, 2.5), "top margin is not 2.5 cm"),
        (near(section.bottom_margin.cm, 2.5), "bottom margin is not 2.5 cm"),
        (near(section.left_margin.cm, 3.0), "left margin is not 3.0 cm"),
        (near(section.right_margin.cm, 2.5), "right margin is not 2.5 cm"),
    ]
    for ok, message in checks:
        if not ok:
            fail(message)

    body = doc.styles["Thesis Body"]
    body_pf = body.paragraph_format
    if body_pf.line_spacing_rule != WD_LINE_SPACING.EXACTLY or not near(body_pf.line_spacing.pt, 20.0, 0.1):
        fail("body line spacing is not exact 20 pt")
    if not near(body_pf.first_line_indent.pt, 24.0, 0.1):
        fail("body first-line indent is not 24 pt")
    if style_east_asia_font(body) != "宋体" or style_ascii_font(body) != "Times New Roman":
        fail("body font is not 宋体 / Times New Roman")

    heading1 = doc.styles["Heading 1"]
    heading1_pf = heading1.paragraph_format
    if heading1_pf.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        fail("Heading 1 is not centered")
    if heading1_pf.line_spacing_rule != WD_LINE_SPACING.EXACTLY or not near(heading1_pf.line_spacing.pt, 24.0, 0.1):
        fail("Heading 1 line spacing is not exact 24 pt")


def validate_chapters(doc: Document, book_root: Path):
    chapter_paths = sorted((book_root / "chapters").glob("chapter-*.md"))
    if len(chapter_paths) != 96:
        fail(f"expected 96 chapter files, found {len(chapter_paths)}")

    expected_titles: list[str] = []
    expected_body = Counter()
    for chapter_path in chapter_paths:
        title, body = chapter_title_and_body(chapter_path)
        expected_titles.append(title)
        expected_body.update(body)

    heading1_paragraphs = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.style is not None and paragraph.style.name == "Heading 1" and paragraph.text.strip()
    ]
    heading1_texts = [paragraph.text.strip() for paragraph in heading1_paragraphs]
    chapter_headings = [text for text in heading1_texts if text in set(expected_titles)]
    if chapter_headings != expected_titles:
        missing = [title for title in expected_titles if title not in chapter_headings]
        fail("chapter headings are missing or out of order: " + ", ".join(missing[:5]))

    doubled = [text for text in heading1_texts if DOUBLE_CHAPTER_RE.match(text)]
    if doubled:
        fail("duplicated chapter numbering in headings: " + ", ".join(doubled[:5]))

    expected_title_set = set(expected_titles)
    bad_heading_fonts = [
        paragraph.text.strip()
        for paragraph in heading1_paragraphs
        if paragraph.text.strip() in expected_title_set and not paragraph_has_run_font(paragraph, "黑体", "Times New Roman")
    ]
    if bad_heading_fonts:
        fail("chapter heading direct font is not 黑体 / Times New Roman: " + bad_heading_fonts[0])

    body_paragraphs = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.style is not None and paragraph.style.name == "Thesis Body" and paragraph.text.strip()
    ]
    bad_body_fonts = [
        paragraph.text.strip()
        for paragraph in body_paragraphs
        if not paragraph_has_run_font(paragraph, "宋体", "Times New Roman")
    ]
    if bad_body_fonts:
        fail("body direct font is not 宋体 / Times New Roman: " + bad_body_fonts[0][:80])

    actual_body = Counter(paragraph.text.strip() for paragraph in body_paragraphs)
    missing_body = [text for text, count in expected_body.items() if actual_body[text] < count]
    if missing_body:
        fail("missing chapter body paragraphs: " + missing_body[0][:80])

    return {
        "chapter_files": len(chapter_paths),
        "chapter_headings": len(chapter_headings),
        "chapter_paragraphs": sum(expected_body.values()),
        "h1_total": len(heading1_texts),
        "body_paragraphs": sum(actual_body.values()),
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", default=str(DEFAULT_DOCX))
    parser.add_argument("--book-root", default=str(BOOK_ROOT))
    args = parser.parse_args()

    docx_path = Path(args.docx)
    book_root = Path(args.book_root)
    if not docx_path.exists():
        fail(f"docx does not exist: {docx_path}")

    zip_stats = validate_zip(docx_path)
    doc = Document(docx_path)
    validate_layout_and_styles(doc)
    stats = validate_chapters(doc, book_root)

    print(f"VALID_DOCX={docx_path.resolve()}")
    print(f"UPDATE_FIELDS_ON_OPEN={str(zip_stats['update_fields']).upper()}")
    for key, value in stats.items():
        print(f"{key.upper()}={value}")


if __name__ == "__main__":
    main()
